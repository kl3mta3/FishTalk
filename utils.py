"""
FishTalk — Utility functions.

Handles ffmpeg detection, file reading (txt/pdf/docx), audio export,
document export, and system monitoring helpers.
"""

import os
import sys
import shutil
import logging

import psutil
import pdfplumber
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

logger = logging.getLogger("FishTalk.utils")

# ---------------------------------------------------------------------------
# FFmpeg setup
# ---------------------------------------------------------------------------

def get_app_dir() -> str:
    """Return the directory where the application lives."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _download_ffmpeg(bin_dir: str, on_progress=None) -> bool:
    """
    Download a pre-built FFmpeg Windows binary from gyan.dev and extract
    ffmpeg.exe + ffprobe.exe into bin_dir.

    Returns True on success, False on failure.
    """
    import urllib.request
    import zipfile
    import tempfile

    # Essentials build: ~75 MB, contains ffmpeg.exe and ffprobe.exe
    URL = "https://www.gyan.dev/ffmpeg/builds/ffmpeg-release-essentials.zip"

    os.makedirs(bin_dir, exist_ok=True)

    try:
        if on_progress:
            on_progress("Downloading FFmpeg...", 0.05)

        logger.info("Downloading FFmpeg from %s", URL)

        tmp_zip = os.path.join(tempfile.gettempdir(), "ffmpeg_download.zip")

        def _reporthook(block_num, block_size, total_size):
            if on_progress and total_size > 0:
                frac = min(block_num * block_size / total_size, 1.0)
                on_progress(f"Downloading FFmpeg... {int(frac * 100)}%", frac * 0.8)

        urllib.request.urlretrieve(URL, tmp_zip, reporthook=_reporthook)

        if on_progress:
            on_progress("Extracting FFmpeg...", 0.85)

        with zipfile.ZipFile(tmp_zip, "r") as zf:
            for member in zf.namelist():
                filename = os.path.basename(member)
                if filename in ("ffmpeg.exe", "ffprobe.exe"):
                    source = zf.open(member)
                    dest_path = os.path.join(bin_dir, filename)
                    with open(dest_path, "wb") as dest:
                        dest.write(source.read())
                    logger.info("Extracted %s -> %s", filename, dest_path)

        try:
            os.remove(tmp_zip)
        except OSError:
            pass

        if on_progress:
            on_progress("FFmpeg ready", 1.0)

        return os.path.isfile(os.path.join(bin_dir, "ffmpeg.exe"))

    except Exception as exc:
        logger.error("FFmpeg auto-download failed: %s", exc)
        return False


def setup_ffmpeg(on_progress=None) -> bool:
    """
    Locate FFmpeg and ensure it is available on the OS PATH so libraries
    like pydub can find it. If FFmpeg is not present, automatically
    downloads the official Windows build from gyan.dev into bin/.

    Accepts an optional on_progress(message, fraction) callback for
    splash screen integration.
    """
    app_dir = get_app_dir()
    bin_dir = os.path.join(app_dir, "bin")
    bundled_ffmpeg = os.path.join(bin_dir, "ffmpeg.exe")

    def _activate(ffmpeg_path: str) -> bool:
        """Put the containing folder on PATH and configure pydub."""
        folder = os.path.dirname(ffmpeg_path)
        current_path = os.environ.get("PATH", "")
        if folder not in current_path:
            os.environ["PATH"] = folder + os.pathsep + current_path
        try:
            from pydub import AudioSegment
            AudioSegment.converter = ffmpeg_path
        except ImportError:
            pass
        return True

    # 1. Bundled bin/ folder
    if os.path.isfile(bundled_ffmpeg):
        logger.info("FFmpeg found (bundled): %s", bundled_ffmpeg)
        return _activate(bundled_ffmpeg)

    # 2. System PATH
    path_ffmpeg = shutil.which("ffmpeg")
    if path_ffmpeg:
        logger.info("FFmpeg found (PATH): %s", path_ffmpeg)
        return _activate(path_ffmpeg)

    # 3. Auto-download
    logger.warning("FFmpeg not found -- attempting auto-download.")
    success = _download_ffmpeg(bin_dir, on_progress=on_progress)
    if success:
        logger.info("FFmpeg auto-download succeeded.")
        return _activate(bundled_ffmpeg)

    logger.error("FFmpeg could not be installed automatically.")
    return False




def is_ffmpeg_available() -> bool:
    """Quick check without mutating state."""
    bundled = os.path.join(get_app_dir(), "bin", "ffmpeg.exe")
    return os.path.isfile(bundled) or shutil.which("ffmpeg") is not None


# ---------------------------------------------------------------------------
# Fish-Speech auto-setup
# ---------------------------------------------------------------------------

def setup_fish_speech(dest_dir: str, on_progress=None) -> bool:
    """
    Ensure Fish-Speech v1.4.3 code and model checkpoints are present.

    If the dest_dir does not exist or is empty:
      1. Downloads the v1.4.3 source zip from GitHub (~20 MB).
      2. Downloads the checkpoint files from HuggingFace
         (fishaudio/fish-speech-1.4, ~1.5 GB).

    on_progress(message, fraction) is called throughout so the splash screen
    can show download progress.

    Returns True if fish-speech is ready, False on failure.
    """
    import urllib.request
    import zipfile
    import tempfile

    CODE_URL = "https://github.com/fishaudio/fish-speech/archive/refs/tags/v1.4.3.zip"

    checkpoints_base = os.path.join(dest_dir, "checkpoints")

    # --- Step 1: Fish-Speech code ---
    if not os.path.isdir(dest_dir) or not os.listdir(dest_dir):
        try:
            if on_progress:
                on_progress("Downloading Fish-Speech code (~20 MB)...", 0.05)

            logger.info("Downloading Fish-Speech source from %s", CODE_URL)
            tmp_zip = os.path.join(tempfile.gettempdir(), "fish_speech_src.zip")

            def _hook(b, bs, total):
                if on_progress and total > 0:
                    frac = min(b * bs / total, 1.0) * 0.2
                    on_progress(f"Downloading Fish-Speech code... {int(frac * 5 * 100)}%", frac)

            urllib.request.urlretrieve(CODE_URL, tmp_zip, reporthook=_hook)

            if on_progress:
                on_progress("Extracting Fish-Speech code...", 0.22)

            os.makedirs(dest_dir, exist_ok=True)
            with zipfile.ZipFile(tmp_zip, "r") as zf:
                for member in zf.infolist():
                    # Strip the top-level "fish-speech-1.4.3/" prefix
                    parts = member.filename.split("/", 1)
                    if len(parts) < 2 or not parts[1]:
                        continue
                    target = os.path.join(dest_dir, parts[1])
                    if member.is_dir():
                        os.makedirs(target, exist_ok=True)
                    else:
                        os.makedirs(os.path.dirname(target), exist_ok=True)
                        with zf.open(member) as src, open(target, "wb") as dst:
                            dst.write(src.read())

            try:
                os.remove(tmp_zip)
            except OSError:
                pass

            logger.info("Fish-Speech code extracted to %s", dest_dir)

        except Exception as exc:
            logger.error("Fish-Speech code download failed: %s", exc)
            return False

    # --- Step 2: Model checkpoints ---
    checkpoint_sets = [
        ("fish-speech-1.4", "fishaudio/fish-speech-1.4"),
        ("fish-speech-1.5", "fishaudio/fish-speech-1.5"),
    ]

    for ckpt_name, hf_repo in checkpoint_sets:
        ckpt_dir = os.path.join(dest_dir, "checkpoints", ckpt_name)
        has_weights = (
            os.path.isdir(ckpt_dir) and any(
                f.endswith((".pth", ".bin", ".safetensors"))
                for f in os.listdir(ckpt_dir)
            )
        ) if os.path.isdir(ckpt_dir) else False

        if not has_weights:
            try:
                from huggingface_hub import snapshot_download
            except ImportError:
                logger.error("huggingface_hub not installed; cannot download checkpoints.")
                return False

            if on_progress:
                on_progress(f"Downloading {ckpt_name} checkpoints (~1.5 GB)...", 0.25)

            logger.info("Downloading checkpoints from HuggingFace: %s", hf_repo)
            try:
                snapshot_download(
                    repo_id=hf_repo,
                    local_dir=ckpt_dir,
                    ignore_patterns=["*.md", "*.txt"],
                )
                logger.info("Checkpoints downloaded to %s", ckpt_dir)
            except Exception as exc:
                logger.error("Checkpoint download failed for %s: %s", ckpt_name, exc)
                return False

    if on_progress:
        on_progress("Fish-Speech ready", 1.0)

    return True




# ---------------------------------------------------------------------------
# File readers
# ---------------------------------------------------------------------------

def read_txt(path: str) -> str:
    """Read a plain text file."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_pdf(path: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def read_docx(path: str) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_epub(path: str) -> str:
    """
    Extract text from an EPUB file.

    Reads chapters in spine order, strips HTML tags, and joins paragraphs
    with double newlines so the sentence splitter sees proper boundaries.
    """
    import html
    import re
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        """Minimal HTML-to-text converter that preserves paragraph breaks."""
        BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
                      "li", "tr", "br", "blockquote", "section", "article"}

        def __init__(self):
            super().__init__()
            self._parts: list = []
            self._current = []
            self._skip = False

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style", "head"):
                self._skip = True
            if tag in self.BLOCK_TAGS and self._current:
                self._flush()

        def handle_endtag(self, tag):
            if tag in ("script", "style", "head"):
                self._skip = False
                self._current.clear()
            if tag in self.BLOCK_TAGS:
                self._flush()

        def handle_data(self, data):
            if not self._skip:
                data = data.strip()
                if data:
                    self._current.append(data)

        def handle_entityref(self, name):
            self._current.append(html.unescape(f"&{name};"))

        def handle_charref(self, name):
            self._current.append(html.unescape(f"&#{name};"))

        def _flush(self):
            text = " ".join(self._current).strip()
            if text:
                self._parts.append(text)
            self._current.clear()

        def get_text(self) -> str:
            self._flush()
            return "\n\n".join(self._parts)

    try:
        import ebooklib
        from ebooklib import epub
    except ImportError:
        raise ImportError(
            "ebooklib is required for EPUB support.\n"
            "Install it with: pip install ebooklib"
        )

    book = epub.read_epub(path, options={"ignore_ncx": True})
    chapters: list = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        content = item.get_content()
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        extractor = _TextExtractor()
        extractor.feed(content)
        text = extractor.get_text().strip()
        if text:
            chapters.append(text)

    if not chapters:
        raise ValueError("No readable text found in EPUB.")

    return "\n\n".join(chapters)


def read_file(path: str) -> str:

    """
    Dispatch to the correct reader based on file extension.
    Supported: .txt, .pdf, .docx, .epub
    """
    ext = os.path.splitext(path)[1].lower()
    readers = {
        ".txt":  read_txt,
        ".pdf":  read_pdf,
        ".docx": read_docx,
        ".epub": read_epub,
    }
    reader = readers.get(ext)
    if reader is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return reader(path)



# ---------------------------------------------------------------------------
# Audio export
# ---------------------------------------------------------------------------

def export_mp3(wav_path: str, out_path: str) -> str:
    """Convert a WAV file to MP3 using pydub + ffmpeg."""
    from pydub import AudioSegment

    if not is_ffmpeg_available():
        raise RuntimeError("ffmpeg is not available. Cannot export to MP3.")
    setup_ffmpeg()
    audio = AudioSegment.from_wav(wav_path)
    audio.export(out_path, format="mp3", bitrate="192k")
    logger.info("Exported MP3: %s", out_path)
    return out_path


def apply_speed(wav_path: str, speed: float, out_path: str = None) -> str:
    """
    Adjust playback speed of a WAV file.
    speed < 1.0 = slower, speed > 1.0 = faster.
    """
    from pydub import AudioSegment

    if abs(speed - 1.0) < 0.01:
        return wav_path

    if out_path is None:
        base, ext = os.path.splitext(wav_path)
        out_path = f"{base}_speed{ext}"

    audio = AudioSegment.from_wav(wav_path)
    # Change frame rate for speed adjustment, then restore sample rate
    adjusted = audio._spawn(
        audio.raw_data,
        overrides={"frame_rate": int(audio.frame_rate * speed)}
    ).set_frame_rate(audio.frame_rate)
    adjusted.export(out_path, format="wav")
    return out_path


def apply_volume(wav_path: str, volume_db: float, out_path: str = None) -> str:
    """Adjust volume of a WAV file by the given dB amount."""
    from pydub import AudioSegment

    if abs(volume_db) < 0.1:
        return wav_path

    if out_path is None:
        base, ext = os.path.splitext(wav_path)
        out_path = f"{base}_vol{ext}"

    audio = AudioSegment.from_wav(wav_path)
    adjusted = audio + volume_db
    adjusted.export(out_path, format="wav")
    return out_path


# ---------------------------------------------------------------------------
# Document export
# ---------------------------------------------------------------------------

def export_txt(text: str, path: str) -> str:
    """Save text to a .txt file."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    logger.info("Exported TXT: %s", path)
    return path


def export_docx(text: str, path: str) -> str:
    """Save text to a .docx file."""
    doc = DocxDocument()
    for para in text.split("\n"):
        doc.add_paragraph(para)
    doc.save(path)
    logger.info("Exported DOCX: %s", path)
    return path


def export_pdf(text: str, path: str) -> str:
    """Save text to a .pdf file using ReportLab."""
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for para in text.split("\n"):
        if para.strip():
            story.append(Paragraph(para, styles["Normal"]))
            story.append(Spacer(1, 6))
    if not story:
        story.append(Paragraph("(empty)", styles["Normal"]))
    doc.build(story)
    logger.info("Exported PDF: %s", path)
    return path


# ---------------------------------------------------------------------------
# System monitoring
# ---------------------------------------------------------------------------

def get_ram_usage() -> dict:
    """
    Return RAM usage info for the current process and system.

    Returns dict with keys:
      - process_mb: float — RAM used by this process
      - system_percent: float — system-wide RAM usage percentage
      - system_used_gb: float — system RAM used in GB
      - system_total_gb: float — total system RAM in GB
    """
    proc = psutil.Process(os.getpid())
    mem_info = proc.memory_info()
    sys_mem = psutil.virtual_memory()
    return {
        "process_mb": mem_info.rss / (1024 * 1024),
        "system_percent": sys_mem.percent,
        "system_used_gb": sys_mem.used / (1024 ** 3),
        "system_total_gb": sys_mem.total / (1024 ** 3),
    }
